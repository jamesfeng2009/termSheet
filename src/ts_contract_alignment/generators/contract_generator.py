"""Contract generator implementation for the TS Contract Alignment System."""

import copy
import logging
import uuid
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from ..interfaces.generator import (
    GeneratedContract,
    IContractGenerator,
    Modification,
)
from ..models.alignment import AlignmentMatch, AlignmentResult
from ..models.document import DocumentSection, ParsedDocument, TextSegment
from ..models.enums import ActionType
from ..models.extraction import ExtractedTerm, TSExtractionResult
from ..models.template import AnalyzedClause, FillableSegment, TemplateAnalysisResult
from .conflict_handler import ConflictHandler, ConflictHandlerConfig, ConflictType


logger = logging.getLogger(__name__)


@dataclass
class ConflictRecord:
    """Record of a formatting conflict during generation."""
    modification_id: str
    conflict_type: str
    original_format: Dict[str, Any]
    attempted_format: Dict[str, Any]
    resolution: str
    timestamp: str = field(default_factory=lambda: datetime.utcnow().isoformat())


@dataclass
class AnnotationConfig:
    """Configuration for annotation display."""
    show_source_id: bool = True
    show_action_type: bool = True
    show_confidence: bool = True
    annotation_style: str = "inline"  # "inline" or "margin"
    highlight_insertions: bool = True
    highlight_overrides: bool = True
    insertion_color: str = "green"
    override_color: str = "yellow"


class ContractGenerator(IContractGenerator):
    """
    Contract generator that applies alignment results to templates.
    
    Implements value insertion, override, annotation, and export
    functionality while preserving document formatting.
    """

    # Color mapping for highlights
    COLOR_MAP = {
        "green": WD_COLOR_INDEX.BRIGHT_GREEN,
        "yellow": WD_COLOR_INDEX.YELLOW,
        "red": WD_COLOR_INDEX.RED,
        "blue": WD_COLOR_INDEX.TURQUOISE,
    }

    def __init__(
        self,
        output_dir: str = "data/generated",
        annotation_config: Optional[AnnotationConfig] = None,
        conflict_config: Optional[ConflictHandlerConfig] = None,
    ):
        """
        Initialize the contract generator.
        
        Args:
            output_dir: Directory for generated contract files.
            annotation_config: Configuration for annotations.
            conflict_config: Configuration for conflict handling.
        """
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.annotation_config = annotation_config or AnnotationConfig()
        self.conflict_handler = ConflictHandler(conflict_config)
        self.conflicts: List[ConflictRecord] = []

    def generate(
        self,
        template_doc: ParsedDocument,
        alignment_result: AlignmentResult,
        ts_result: TSExtractionResult,
    ) -> GeneratedContract:
        """
        Generate a completed contract by applying alignment mappings.
        
        Args:
            template_doc: The parsed contract template document.
            alignment_result: The alignment result with term-to-clause mappings.
            ts_result: The extracted TS terms.
            
        Returns:
            GeneratedContract with all modifications applied.
        """
        self.conflicts = []  # Reset conflicts for new generation
        self.conflict_handler.clear_conflicts()  # Reset conflict handler
        
        contract_id = str(uuid.uuid4())
        modifications: List[Modification] = []
        
        # Build lookup maps for efficient access
        terms_by_id = {term.id: term for term in ts_result.terms}
        
        # Process each alignment match
        for match in alignment_result.matches:
            if match.action == ActionType.SKIP:
                continue
                
            term = terms_by_id.get(match.ts_term_id)
            if not term:
                logger.warning(f"Term not found for match: {match.ts_term_id}")
                continue
            
            modification = self._create_modification(match, term, template_doc)
            if modification:
                modifications.append(modification)
        
        # Generate file paths
        timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        base_name = f"contract_{contract_id[:8]}_{timestamp}"
        revision_path = str(self.output_dir / f"{base_name}_tracked.docx")
        clean_path = str(self.output_dir / f"{base_name}_clean.docx")
        
        return GeneratedContract(
            id=contract_id,
            template_document_id=template_doc.id,
            ts_document_id=ts_result.document_id,
            modifications=modifications,
            revision_tracked_path=revision_path,
            clean_version_path=clean_path,
            generation_timestamp=datetime.utcnow().isoformat(),
        )

    def _create_modification(
        self,
        match: AlignmentMatch,
        term: ExtractedTerm,
        template_doc: ParsedDocument,
    ) -> Optional[Modification]:
        """
        Create a modification record for a single alignment match.
        
        Args:
            match: The alignment match.
            term: The extracted TS term.
            template_doc: The template document.
            
        Returns:
            Modification record or None if creation fails.
        """
        # Find the target location in the template
        original_text, location_start, location_end = self._find_target_location(
            match, template_doc
        )
        
        # Determine the new text value
        new_text = self._format_term_value(term)
        
        # Build annotations
        annotations = self._build_annotations(match, term)
        
        return Modification(
            id=str(uuid.uuid4()),
            match_id=match.id,
            original_text=original_text,
            new_text=new_text,
            location_start=location_start,
            location_end=location_end,
            action=match.action,
            source_ts_paragraph_id=term.source_paragraph_id,
            confidence=match.confidence,
            annotations=annotations,
        )

    def _find_target_location(
        self,
        match: AlignmentMatch,
        template_doc: ParsedDocument,
    ) -> Tuple[str, int, int]:
        """
        Find the target location in the template for a modification.
        
        Args:
            match: The alignment match.
            template_doc: The template document.
            
        Returns:
            Tuple of (original_text, start_position, end_position).
        """
        # Search through sections to find the matching clause
        for section in template_doc.sections:
            result = self._search_section_for_clause(section, match.clause_id)
            if result:
                return result
            
            # Search children recursively
            for child in section.children:
                result = self._search_section_recursive(child, match.clause_id)
                if result:
                    return result
        
        # Default: return empty location if not found
        return "", 0, 0

    def _search_section_for_clause(
        self,
        section: DocumentSection,
        clause_id: str,
    ) -> Optional[Tuple[str, int, int]]:
        """Search a section for a clause ID."""
        if section.id == clause_id:
            # Found the section - return its content
            if section.segments:
                first_seg = section.segments[0]
                last_seg = section.segments[-1]
                content = " ".join(seg.content for seg in section.segments)
                return content, first_seg.start_pos, last_seg.end_pos
            return section.title or "", 0, 0
        return None

    def _search_section_recursive(
        self,
        section: DocumentSection,
        clause_id: str,
    ) -> Optional[Tuple[str, int, int]]:
        """Recursively search sections for a clause ID."""
        result = self._search_section_for_clause(section, clause_id)
        if result:
            return result
        
        for child in section.children:
            result = self._search_section_recursive(child, clause_id)
            if result:
                return result
        
        return None

    def _format_term_value(self, term: ExtractedTerm) -> str:
        """
        Format a term value for insertion into the contract.
        
        Args:
            term: The extracted term.
            
        Returns:
            Formatted string value.
        """
        value = term.value
        
        if value is None:
            return ""
        
        if isinstance(value, dict):
            # Handle structured values
            return self._format_structured_value(value, term.category)
        
        if isinstance(value, (int, float)):
            # Format numbers appropriately
            return self._format_number(value, term.category)
        
        return str(value)

    def _format_structured_value(self, value: Dict[str, Any], category: Any) -> str:
        """Format a structured value dictionary."""
        parts = []
        for key, val in value.items():
            if val is not None:
                parts.append(f"{key}: {val}")
        return "; ".join(parts)

    def _format_number(self, value: float, category: Any) -> str:
        """Format a numeric value based on category."""
        from ..models.enums import TermCategory
        
        if category == TermCategory.INVESTMENT_AMOUNT:
            return f"USD {value:,.2f}"
        elif category == TermCategory.VALUATION:
            return f"USD {value:,.2f}"
        elif category in (TermCategory.LIQUIDATION_PREFERENCE, TermCategory.ANTI_DILUTION):
            return f"{value:.2f}x"
        else:
            return f"{value:,.2f}"

    def _build_annotations(
        self,
        match: AlignmentMatch,
        term: ExtractedTerm,
    ) -> Dict[str, Any]:
        """
        Build annotation dictionary for a modification.
        
        Args:
            match: The alignment match.
            term: The extracted term.
            
        Returns:
            Dictionary of annotation data.
        """
        annotations = {
            "timestamp": datetime.utcnow().isoformat(),
        }
        
        if self.annotation_config.show_source_id:
            annotations["source_ts_paragraph_id"] = term.source_paragraph_id
            annotations["source_section_id"] = term.source_section_id
        
        if self.annotation_config.show_action_type:
            annotations["action_type"] = match.action.value
        
        if self.annotation_config.show_confidence:
            annotations["confidence_score"] = match.confidence
        
        annotations["match_method"] = match.match_method.value
        annotations["term_category"] = term.category.value
        annotations["needs_review"] = match.needs_review
        
        return annotations

    def apply_modifications_to_document(
        self,
        source_path: str,
        modifications: List[Modification],
        with_annotations: bool = True,
    ) -> Document:
        """
        Apply modifications to a Word document.
        
        Args:
            source_path: Path to the source template document.
            modifications: List of modifications to apply.
            with_annotations: Whether to add annotations.
            
        Returns:
            Modified Document object.
        """
        doc = Document(source_path)
        
        # Sort modifications by location (reverse order to preserve positions)
        sorted_mods = sorted(
            modifications,
            key=lambda m: m.location_start,
            reverse=True,
        )
        
        for mod in sorted_mods:
            self._apply_single_modification(doc, mod, with_annotations)
        
        return doc

    def _apply_single_modification(
        self,
        doc: Document,
        mod: Modification,
        with_annotations: bool,
    ) -> None:
        """
        Apply a single modification to the document.
        
        Args:
            doc: The Document object.
            mod: The modification to apply.
            with_annotations: Whether to add annotations.
        """
        # Find the paragraph containing the modification
        for para in doc.paragraphs:
            if mod.original_text and mod.original_text in para.text:
                try:
                    self._modify_paragraph(para, mod, with_annotations)
                    return
                except Exception as e:
                    # Record conflict and preserve original
                    self._record_conflict(mod, "modification_failed", str(e))
                    return
        
        # If original text not found, try to append to appropriate section
        if mod.action == ActionType.INSERT:
            self._insert_new_content(doc, mod, with_annotations)

    def _modify_paragraph(
        self,
        para: Paragraph,
        mod: Modification,
        with_annotations: bool,
    ) -> None:
        """
        Modify a paragraph with the new value.
        
        Args:
            para: The paragraph to modify.
            mod: The modification details.
            with_annotations: Whether to add annotations.
        """
        original_formatting = self._capture_formatting(para)
        
        if mod.action == ActionType.OVERRIDE:
            # Replace the original text with new text
            new_text = para.text.replace(mod.original_text, mod.new_text)
            self._set_paragraph_text_with_formatting(
                para, new_text, original_formatting, mod, with_annotations
            )
        elif mod.action == ActionType.INSERT:
            # Append the new text
            self._append_to_paragraph(para, mod, with_annotations)

    def _capture_formatting(self, para: Paragraph) -> Dict[str, Any]:
        """Capture the formatting of a paragraph."""
        formatting = {
            "alignment": para.alignment,
            "style": para.style.name if para.style else None,
        }
        
        if para.runs:
            first_run = para.runs[0]
            formatting["bold"] = first_run.bold
            formatting["italic"] = first_run.italic
            formatting["underline"] = first_run.underline
            formatting["font_name"] = first_run.font.name
            formatting["font_size"] = first_run.font.size
        
        return formatting

    def _set_paragraph_text_with_formatting(
        self,
        para: Paragraph,
        new_text: str,
        original_formatting: Dict[str, Any],
        mod: Modification,
        with_annotations: bool,
    ) -> None:
        """Set paragraph text while preserving formatting."""
        # Clear existing runs
        para.clear()
        
        # Add the new text
        run = para.add_run(new_text)
        
        # Apply original formatting
        try:
            self._apply_formatting_to_run(run, original_formatting)
        except Exception as e:
            self._record_conflict(mod, "formatting_conflict", str(e))
            # Continue with default formatting
        
        # Add highlight for modifications
        if with_annotations:
            self._add_modification_highlight(run, mod)
        
        # Add annotation comment if configured
        if with_annotations and self.annotation_config.annotation_style == "inline":
            self._add_inline_annotation(para, mod)

    def _apply_formatting_to_run(
        self,
        run: Run,
        formatting: Dict[str, Any],
    ) -> None:
        """Apply formatting to a run."""
        if formatting.get("bold") is not None:
            run.bold = formatting["bold"]
        if formatting.get("italic") is not None:
            run.italic = formatting["italic"]
        if formatting.get("underline") is not None:
            run.underline = formatting["underline"]
        if formatting.get("font_name"):
            run.font.name = formatting["font_name"]
        if formatting.get("font_size"):
            run.font.size = formatting["font_size"]

    def _add_modification_highlight(self, run: Run, mod: Modification) -> None:
        """Add highlight color to indicate modification type."""
        if mod.action == ActionType.INSERT and self.annotation_config.highlight_insertions:
            color = self.COLOR_MAP.get(
                self.annotation_config.insertion_color,
                WD_COLOR_INDEX.BRIGHT_GREEN,
            )
            run.font.highlight_color = color
        elif mod.action == ActionType.OVERRIDE and self.annotation_config.highlight_overrides:
            color = self.COLOR_MAP.get(
                self.annotation_config.override_color,
                WD_COLOR_INDEX.YELLOW,
            )
            run.font.highlight_color = color

    def _add_inline_annotation(self, para: Paragraph, mod: Modification) -> None:
        """Add inline annotation text after the modification."""
        annotation_parts = []
        
        if self.annotation_config.show_source_id:
            annotation_parts.append(f"[TS:{mod.source_ts_paragraph_id}]")
        
        if self.annotation_config.show_action_type:
            annotation_parts.append(f"[{mod.action.value.upper()}]")
        
        if self.annotation_config.show_confidence:
            annotation_parts.append(f"[Conf:{mod.confidence:.2f}]")
        
        if annotation_parts:
            annotation_text = " " + " ".join(annotation_parts)
            annotation_run = para.add_run(annotation_text)
            annotation_run.font.size = Pt(8)
            annotation_run.font.color.rgb = RGBColor(128, 128, 128)
            annotation_run.italic = True

    def _append_to_paragraph(
        self,
        para: Paragraph,
        mod: Modification,
        with_annotations: bool,
    ) -> None:
        """Append new content to a paragraph."""
        run = para.add_run(f" {mod.new_text}")
        
        if with_annotations:
            self._add_modification_highlight(run, mod)
            self._add_inline_annotation(para, mod)

    def _insert_new_content(
        self,
        doc: Document,
        mod: Modification,
        with_annotations: bool,
    ) -> None:
        """Insert new content when original location not found."""
        # Add a new paragraph at the end
        para = doc.add_paragraph()
        run = para.add_run(mod.new_text)
        
        if with_annotations:
            self._add_modification_highlight(run, mod)
            self._add_inline_annotation(para, mod)

    def _record_conflict(
        self,
        mod: Modification,
        conflict_type: str,
        details: str,
    ) -> None:
        """Record a formatting conflict."""
        conflict = ConflictRecord(
            modification_id=mod.id,
            conflict_type=conflict_type,
            original_format={},
            attempted_format={},
            resolution=f"Preserved original format. Details: {details}",
        )
        self.conflicts.append(conflict)
        logger.warning(f"Conflict recorded: {conflict_type} for modification {mod.id}")

    def export_docx(
        self,
        contract: GeneratedContract,
        with_revisions: bool = True,
    ) -> str:
        """
        Export the generated contract to a .docx file.
        
        Args:
            contract: The generated contract to export.
            with_revisions: Whether to include revision tracking marks.
            
        Returns:
            File path to the exported .docx document.
        """
        # For now, return the appropriate path based on revision flag
        if with_revisions:
            return contract.revision_tracked_path
        return contract.clean_version_path

    def export_with_template(
        self,
        contract: GeneratedContract,
        template_path: str,
        with_revisions: bool = True,
    ) -> str:
        """
        Export contract by applying modifications to a template file.
        
        Args:
            contract: The generated contract.
            template_path: Path to the template .docx file.
            with_revisions: Whether to include revision marks.
            
        Returns:
            Path to the exported file.
        """
        # Apply modifications to the template
        doc = self.apply_modifications_to_document(
            template_path,
            contract.modifications,
            with_annotations=with_revisions,
        )
        
        # Determine output path
        output_path = (
            contract.revision_tracked_path
            if with_revisions
            else contract.clean_version_path
        )
        
        # Ensure output directory exists
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        
        # Save the document
        doc.save(output_path)
        
        return output_path

    def export_both_versions(
        self,
        contract: GeneratedContract,
        template_path: str,
    ) -> Tuple[str, str]:
        """
        Export both revision-tracked and clean versions.
        
        Args:
            contract: The generated contract.
            template_path: Path to the template .docx file.
            
        Returns:
            Tuple of (revision_tracked_path, clean_version_path).
        """
        # Export revision-tracked version
        revision_path = self.export_with_template(
            contract, template_path, with_revisions=True
        )
        
        # Export clean version
        clean_path = self.export_with_template(
            contract, template_path, with_revisions=False
        )
        
        return revision_path, clean_path

    def get_conflicts(self) -> List[ConflictRecord]:
        """Get all recorded conflicts from the last generation."""
        return self.conflicts.copy()

    def clear_conflicts(self) -> None:
        """Clear recorded conflicts."""
        self.conflicts = []
