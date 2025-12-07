"""Final document export after review."""

import os
from typing import List, Optional
from datetime import datetime
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_COLOR_INDEX

from ..interfaces.review import ReviewSession, ReviewItem, ReviewAction
from ..interfaces.generator import GeneratedContract, Modification
from ..models.document import ParsedDocument


class FinalExporter:
    """
    Exports finalized documents after review.
    
    Applies accepted changes and reverts rejected changes
    to produce the final contract document.
    """

    def __init__(self, output_dir: str = "data"):
        """
        Initialize the final exporter.
        
        Args:
            output_dir: Directory for output files.
        """
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)

    def export_finalized_contract(
        self,
        template_doc: ParsedDocument,
        contract: GeneratedContract,
        session: ReviewSession,
    ) -> str:
        """
        Export the finalized contract after review.
        
        Args:
            template_doc: Original template document.
            contract: Generated contract with modifications.
            session: Review session with user decisions.
            
        Returns:
            Path to the exported .docx file.
        """
        # Filter modifications based on review decisions
        accepted_mods = self._get_accepted_modifications(
            contract.modifications,
            session.items,
        )

        # Create output document
        output_path = os.path.join(
            self.output_dir,
            f"final_contract_{contract.id}_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.docx"
        )

        # Apply accepted modifications to template
        self._apply_modifications_to_document(
            template_doc,
            accepted_mods,
            output_path,
        )

        return output_path

    def export_with_revisions(
        self,
        template_doc: ParsedDocument,
        contract: GeneratedContract,
        session: ReviewSession,
    ) -> str:
        """
        Export contract with revision marks showing changes.
        
        Args:
            template_doc: Original template document.
            contract: Generated contract with modifications.
            session: Review session with user decisions.
            
        Returns:
            Path to the exported .docx file with revisions.
        """
        output_path = os.path.join(
            self.output_dir,
            f"contract_with_revisions_{contract.id}_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.docx"
        )

        # Apply all modifications with revision tracking
        self._apply_modifications_with_tracking(
            template_doc,
            contract.modifications,
            session.items,
            output_path,
        )

        return output_path

    def _get_accepted_modifications(
        self,
        modifications: List[Modification],
        review_items: List[ReviewItem],
    ) -> List[Modification]:
        """
        Filter modifications to only accepted ones.
        
        Args:
            modifications: All modifications.
            review_items: Review items with decisions.
            
        Returns:
            List of accepted modifications.
        """
        # Create mapping of modification ID to review action
        action_map = {
            item.modification_id: item.action
            for item in review_items
        }

        # Filter to accepted modifications
        accepted = []
        for mod in modifications:
            action = action_map.get(mod.id, ReviewAction.PENDING)
            if action == ReviewAction.ACCEPT:
                accepted.append(mod)
            elif action == ReviewAction.MODIFY:
                # Find the modified version
                for item in review_items:
                    if item.modification_id == mod.id:
                        # Create a new modification with updated text
                        modified_mod = Modification(
                            id=mod.id,
                            match_id=mod.match_id,
                            original_text=mod.original_text,
                            new_text=item.new_text,
                            location_start=mod.location_start,
                            location_end=mod.location_end,
                            action=mod.action,
                            source_ts_paragraph_id=mod.source_ts_paragraph_id,
                            confidence=mod.confidence,
                            annotations=mod.annotations,
                        )
                        accepted.append(modified_mod)
                        break

        return accepted

    def _apply_modifications_to_document(
        self,
        template_doc: ParsedDocument,
        modifications: List[Modification],
        output_path: str,
    ) -> None:
        """
        Apply modifications to create final document.
        
        Args:
            template_doc: Original template document.
            modifications: Modifications to apply.
            output_path: Path for output file.
        """
        # Load the original document
        # For simplicity, we'll create a new document
        # In a real implementation, this would load and modify the actual template
        doc = Document()
        
        # Add title
        doc.add_heading('Finalized Contract', 0)
        
        # Build text with modifications applied
        text = template_doc.raw_text
        
        # Sort modifications by location (reverse order to maintain positions)
        sorted_mods = sorted(
            modifications,
            key=lambda m: m.location_start,
            reverse=True,
        )
        
        # Apply modifications
        for mod in sorted_mods:
            start = mod.location_start
            end = mod.location_end
            text = text[:start] + mod.new_text + text[end:]
        
        # Add paragraphs to document
        for paragraph_text in text.split('\n'):
            if paragraph_text.strip():
                doc.add_paragraph(paragraph_text)
        
        # Save document
        doc.save(output_path)

    def _apply_modifications_with_tracking(
        self,
        template_doc: ParsedDocument,
        modifications: List[Modification],
        review_items: List[ReviewItem],
        output_path: str,
    ) -> None:
        """
        Apply modifications with revision tracking marks.
        
        Args:
            template_doc: Original template document.
            modifications: All modifications.
            review_items: Review items with decisions.
            output_path: Path for output file.
        """
        # Create mapping of modification ID to review action
        action_map = {
            item.modification_id: item.action
            for item in review_items
        }

        # Load the original document
        doc = Document()
        
        # Add title
        doc.add_heading('Contract with Revisions', 0)
        
        # Build text with tracked changes
        text = template_doc.raw_text
        
        # Sort modifications by location
        sorted_mods = sorted(
            modifications,
            key=lambda m: m.location_start,
        )
        
        # Create annotated text
        last_pos = 0
        annotated_parts = []
        
        for mod in sorted_mods:
            # Add text before modification
            if mod.location_start > last_pos:
                annotated_parts.append({
                    'text': text[last_pos:mod.location_start],
                    'type': 'normal',
                })
            
            # Add modification with annotation
            action = action_map.get(mod.id, ReviewAction.PENDING)
            annotated_parts.append({
                'text': mod.new_text,
                'type': 'modification',
                'action': action,
                'original': mod.original_text,
                'confidence': mod.confidence,
            })
            
            last_pos = mod.location_end
        
        # Add remaining text
        if last_pos < len(text):
            annotated_parts.append({
                'text': text[last_pos:],
                'type': 'normal',
            })
        
        # Create document with annotations
        for part in annotated_parts:
            p = doc.add_paragraph()
            
            if part['type'] == 'normal':
                p.add_run(part['text'])
            else:
                # Add modification with color coding
                run = p.add_run(part['text'])
                
                if part['action'] == ReviewAction.ACCEPT:
                    run.font.color.rgb = RGBColor(0, 128, 0)  # Green
                    run.font.bold = True
                elif part['action'] == ReviewAction.REJECT:
                    run.font.color.rgb = RGBColor(255, 0, 0)  # Red
                    run.font.strike = True
                else:
                    run.font.color.rgb = RGBColor(255, 165, 0)  # Orange
                
                # Add comment about original text if override
                if part['original']:
                    comment_run = p.add_run(f" [Original: {part['original'][:30]}...]")
                    comment_run.font.size = Pt(8)
                    comment_run.font.color.rgb = RGBColor(128, 128, 128)
        
        # Save document
        doc.save(output_path)

    def export_review_summary(
        self,
        session: ReviewSession,
        output_path: Optional[str] = None,
    ) -> str:
        """
        Export a summary of the review session.
        
        Args:
            session: Review session to summarize.
            output_path: Optional output path. If not provided, generates one.
            
        Returns:
            Path to the exported summary file.
        """
        if output_path is None:
            output_path = os.path.join(
                self.output_dir,
                f"review_summary_{session.id}_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.txt"
            )

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write("Review Session Summary\n")
            f.write("=" * 50 + "\n\n")
            f.write(f"Session ID: {session.id}\n")
            f.write(f"Contract ID: {session.contract_id}\n")
            f.write(f"Timestamp: {session.session_timestamp}\n")
            f.write(f"Total Items: {session.total_count}\n")
            f.write(f"Completed: {session.completed_count}\n\n")
            
            # Count by action
            accepted = sum(1 for item in session.items if item.action == ReviewAction.ACCEPT)
            rejected = sum(1 for item in session.items if item.action == ReviewAction.REJECT)
            modified = sum(1 for item in session.items if item.action == ReviewAction.MODIFY)
            pending = sum(1 for item in session.items if item.action == ReviewAction.PENDING)
            
            f.write(f"Accepted: {accepted}\n")
            f.write(f"Rejected: {rejected}\n")
            f.write(f"Modified: {modified}\n")
            f.write(f"Pending: {pending}\n\n")
            
            f.write("Review Items:\n")
            f.write("-" * 50 + "\n")
            
            for item in session.items:
                f.write(f"\nItem ID: {item.modification_id}\n")
                f.write(f"Action: {item.action.value}\n")
                f.write(f"Confidence: {item.confidence:.2%}\n")
                f.write(f"Original: {item.original_text[:50]}...\n")
                f.write(f"New: {item.new_text[:50]}...\n")
                if item.user_comment:
                    f.write(f"Comment: {item.user_comment}\n")
                f.write("-" * 50 + "\n")

        return output_path
