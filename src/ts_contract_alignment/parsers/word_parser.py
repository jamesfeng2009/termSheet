"""Word document (.docx) parser implementation."""

import re
import uuid
from pathlib import Path
from typing import Optional
from zipfile import BadZipFile

from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from ..models.document import DocumentSection, ParsedDocument, TextSegment
from ..models.enums import DocumentType, HeadingLevel
from .exceptions import DocumentCorruptedError, ParseError, UnsupportedFormatError
from .language_detector import detect_language


class WordDocumentParser:
    """
    Parser for Word (.docx) documents.
    
    Extracts section hierarchy, numbering, paragraph structure,
    and formatting information from Word documents.
    """

    # Mapping of Word heading styles to HeadingLevel
    HEADING_STYLE_MAP = {
        "Title": HeadingLevel.TITLE,
        "Heading 1": HeadingLevel.CHAPTER,
        "Heading 2": HeadingLevel.SECTION,
        "Heading 3": HeadingLevel.SUBSECTION,
        "Heading 4": HeadingLevel.PARAGRAPH,
    }

    # Chinese heading patterns
    CHINESE_HEADING_PATTERNS = [
        (re.compile(r'^第[一二三四五六七八九十百千]+[章篇]'), HeadingLevel.CHAPTER),
        (re.compile(r'^第[一二三四五六七八九十百千]+[条节]'), HeadingLevel.SECTION),
        (re.compile(r'^[一二三四五六七八九十]+[、.]'), HeadingLevel.SECTION),
        (re.compile(r'^\([一二三四五六七八九十]+\)'), HeadingLevel.SUBSECTION),
        (re.compile(r'^[（(][0-9]+[)）]'), HeadingLevel.PARAGRAPH),
    ]

    # Numbered heading patterns
    NUMBERED_HEADING_PATTERNS = [
        (re.compile(r'^(\d+)\s*[.、]?\s*(.*)'), HeadingLevel.CHAPTER),
        (re.compile(r'^(\d+\.\d+)\s*[.、]?\s*(.*)'), HeadingLevel.SECTION),
        (re.compile(r'^(\d+\.\d+\.\d+)\s*[.、]?\s*(.*)'), HeadingLevel.SUBSECTION),
    ]

    def __init__(self):
        self._current_position = 0

    def parse(self, file_path: str) -> ParsedDocument:
        """
        Parse a Word document and return its structured representation.
        
        Args:
            file_path: Path to the .docx file.
            
        Returns:
            ParsedDocument containing the structured content.
            
        Raises:
            FileNotFoundError: If the file does not exist.
            UnsupportedFormatError: If the file is not a .docx file.
            DocumentCorruptedError: If the document is corrupted.
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if path.suffix.lower() != ".docx":
            raise UnsupportedFormatError(
                message=f"Unsupported file format: {path.suffix}",
                file_path=file_path,
                location="file extension"
            )
        
        try:
            doc = Document(file_path)
        except (BadZipFile, PackageNotFoundError) as e:
            raise DocumentCorruptedError(
                message="Document is corrupted or not a valid Word file",
                file_path=file_path,
                location="file header",
                details={"original_error": str(e)}
            )
        except Exception as e:
            raise ParseError(
                message=f"Failed to open document: {str(e)}",
                file_path=file_path,
                details={"original_error": str(e)}
            )

        self._current_position = 0
        
        # Extract all text for raw_text
        raw_text = self._extract_raw_text(doc)
        
        # Parse sections with hierarchy
        sections = self._parse_sections(doc)
        
        # Build metadata
        metadata = self._extract_metadata(doc, path)
        
        return ParsedDocument(
            id=str(uuid.uuid4()),
            filename=path.name,
            doc_type=DocumentType.WORD,
            sections=sections,
            metadata=metadata,
            raw_text=raw_text
        )

    def _extract_raw_text(self, doc: Document) -> str:
        """Extract all text content from the document."""
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        return "\n".join(paragraphs)

    def _extract_metadata(self, doc: Document, path: Path) -> dict:
        """Extract document metadata."""
        core_props = doc.core_properties
        
        metadata = {
            "page_count": self._estimate_page_count(doc),
            "word_count": self._count_words(doc),
            "paragraph_count": len(doc.paragraphs),
            "file_size": path.stat().st_size if path.exists() else 0,
        }
        
        # Add core properties if available
        if core_props.title:
            metadata["title"] = core_props.title
        if core_props.author:
            metadata["author"] = core_props.author
        if core_props.created:
            metadata["created"] = core_props.created.isoformat()
        if core_props.modified:
            metadata["modified"] = core_props.modified.isoformat()
            
        return metadata

    def _estimate_page_count(self, doc: Document) -> int:
        """Estimate page count based on content."""
        # Rough estimate: ~500 words per page
        word_count = self._count_words(doc)
        return max(1, word_count // 500 + 1)

    def _count_words(self, doc: Document) -> int:
        """Count total words in the document."""
        total = 0
        for para in doc.paragraphs:
            # Count both English words and Chinese characters
            text = para.text
            # English words
            english_words = len(re.findall(r'[a-zA-Z]+', text))
            # Chinese characters (each counts as a word)
            chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', text))
            total += english_words + chinese_chars
        return total

    def _parse_sections(self, doc: Document) -> list[DocumentSection]:
        """Parse document into hierarchical sections."""
        sections = []
        current_section_stack: list[DocumentSection] = []
        pending_segments: list[TextSegment] = []
        
        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            
            heading_level = self._detect_heading_level(para)
            
            if heading_level is not None:
                # Save pending segments to current section
                if pending_segments and current_section_stack:
                    current_section_stack[-1].segments.extend(pending_segments)
                    pending_segments = []
                elif pending_segments and not current_section_stack:
                    # Create a root section for orphan content
                    root_section = self._create_section(
                        title=None,
                        number=None,
                        level=HeadingLevel.PARAGRAPH,
                        segments=pending_segments
                    )
                    sections.append(root_section)
                    pending_segments = []
                
                # Create new section
                number, title = self._extract_number_and_title(para.text)
                new_section = self._create_section(
                    title=title or para.text.strip(),
                    number=number,
                    level=heading_level
                )
                
                # Find appropriate parent
                while current_section_stack and current_section_stack[-1].level.value >= heading_level.value:
                    current_section_stack.pop()
                
                if current_section_stack:
                    new_section.parent_id = current_section_stack[-1].id
                    current_section_stack[-1].children.append(new_section)
                else:
                    sections.append(new_section)
                
                current_section_stack.append(new_section)
            else:
                # Regular paragraph - create text segment
                segment = self._create_text_segment(para)
                pending_segments.append(segment)
        
        # Handle remaining segments
        if pending_segments:
            if current_section_stack:
                current_section_stack[-1].segments.extend(pending_segments)
            elif pending_segments:
                root_section = self._create_section(
                    title=None,
                    number=None,
                    level=HeadingLevel.PARAGRAPH,
                    segments=pending_segments
                )
                sections.append(root_section)
        
        return sections

    def _detect_heading_level(self, para: Paragraph) -> Optional[HeadingLevel]:
        """Detect if a paragraph is a heading and its level."""
        # Check Word style
        style_name = para.style.name if para.style else None
        if style_name in self.HEADING_STYLE_MAP:
            return self.HEADING_STYLE_MAP[style_name]
        
        text = para.text.strip()
        if not text:
            return None
        
        # Check Chinese heading patterns
        for pattern, level in self.CHINESE_HEADING_PATTERNS:
            if pattern.match(text):
                return level
        
        # Check numbered heading patterns
        for pattern, level in self.NUMBERED_HEADING_PATTERNS:
            if pattern.match(text):
                return level
        
        # Check if paragraph has heading-like formatting
        if self._has_heading_formatting(para):
            return HeadingLevel.SECTION
        
        return None

    def _has_heading_formatting(self, para: Paragraph) -> bool:
        """Check if paragraph has heading-like formatting (bold, larger font)."""
        if not para.runs:
            return False
        
        # Check if entire paragraph is bold
        all_bold = all(run.bold for run in para.runs if run.text.strip())
        
        # Check font size (headings typically > 12pt)
        has_large_font = False
        for run in para.runs:
            if run.font.size and run.font.size.pt > 12:
                has_large_font = True
                break
        
        # Short text that's bold or has large font is likely a heading
        text = para.text.strip()
        is_short = len(text) < 100
        
        return is_short and (all_bold or has_large_font)

    def _extract_number_and_title(self, text: str) -> tuple[Optional[str], Optional[str]]:
        """Extract section number and title from heading text."""
        text = text.strip()
        
        # Try numbered patterns first
        for pattern, _ in self.NUMBERED_HEADING_PATTERNS:
            match = pattern.match(text)
            if match:
                number = match.group(1)
                title = match.group(2).strip() if len(match.groups()) > 1 else None
                return number, title
        
        # Try Chinese patterns
        for pattern, _ in self.CHINESE_HEADING_PATTERNS:
            match = pattern.match(text)
            if match:
                number = match.group(0)
                title = text[len(number):].strip()
                return number, title or None
        
        return None, text

    def _create_section(
        self,
        title: Optional[str],
        number: Optional[str],
        level: HeadingLevel,
        segments: Optional[list[TextSegment]] = None
    ) -> DocumentSection:
        """Create a new DocumentSection."""
        return DocumentSection(
            id=str(uuid.uuid4()),
            title=title,
            number=number,
            level=level,
            segments=segments or [],
            children=[],
            parent_id=None
        )

    def _create_text_segment(self, para: Paragraph) -> TextSegment:
        """Create a TextSegment from a paragraph."""
        text = para.text
        start_pos = self._current_position
        end_pos = start_pos + len(text)
        self._current_position = end_pos + 1  # +1 for newline
        
        formatting = self._extract_formatting(para)
        language = detect_language(text)
        
        return TextSegment(
            id=str(uuid.uuid4()),
            content=text,
            start_pos=start_pos,
            end_pos=end_pos,
            language=language,
            formatting=formatting
        )

    def _extract_formatting(self, para: Paragraph) -> dict:
        """Extract formatting information from a paragraph."""
        formatting = {
            "bold": False,
            "italic": False,
            "underline": False,
            "font_size": 12,
            "font_name": None,
            "alignment": str(para.alignment) if para.alignment else "LEFT",
        }
        
        if para.runs:
            # Use first run's formatting as representative
            first_run = para.runs[0]
            formatting["bold"] = bool(first_run.bold)
            formatting["italic"] = bool(first_run.italic)
            formatting["underline"] = bool(first_run.underline)
            
            if first_run.font.size:
                formatting["font_size"] = first_run.font.size.pt
            if first_run.font.name:
                formatting["font_name"] = first_run.font.name
        
        return formatting
