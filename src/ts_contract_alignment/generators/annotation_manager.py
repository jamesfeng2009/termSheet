"""Annotation manager for contract modifications."""

import logging
from dataclasses import dataclass, field
from datetime import datetime
from enum import Enum
from typing import Any, Dict, List, Optional

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from ..interfaces.generator import Modification
from ..models.enums import ActionType


logger = logging.getLogger(__name__)


class AnnotationStyle(Enum):
    """Styles for displaying annotations."""
    INLINE = "inline"
    MARGIN = "margin"
    COMMENT = "comment"
    HIDDEN = "hidden"


@dataclass
class AnnotationConfig:
    """Configuration for annotation display."""
    show_source_id: bool = True
    show_action_type: bool = True
    show_confidence: bool = True
    style: AnnotationStyle = AnnotationStyle.INLINE
    highlight_insertions: bool = True
    highlight_overrides: bool = True
    insertion_color: str = "green"
    override_color: str = "yellow"
    conflict_color: str = "red"
    font_size: int = 8
    annotation_prefix: str = "【"
    annotation_suffix: str = "】"


@dataclass
class Annotation:
    """Represents an annotation on a modification."""
    id: str
    modification_id: str
    source_ts_paragraph_id: str
    action_type: ActionType
    confidence: float
    timestamp: str
    additional_info: Dict[str, Any] = field(default_factory=dict)

    def to_text(self, config: AnnotationConfig) -> str:
        """Convert annotation to display text."""
        parts = []
        
        if config.show_source_id:
            parts.append(f"TS:{self.source_ts_paragraph_id}")
        
        if config.show_action_type:
            parts.append(self.action_type.value.upper())
        
        if config.show_confidence:
            parts.append(f"置信度:{self.confidence:.0%}")
        
        if not parts:
            return ""
        
        return f"{config.annotation_prefix}{' | '.join(parts)}{config.annotation_suffix}"


class AnnotationManager:
    """
    Manages annotations for contract modifications.
    
    Supports inline annotations, margin comments, and Word comments.
    """

    # Highlight color mapping
    HIGHLIGHT_COLORS = {
        "green": 4,   # Bright Green
        "yellow": 7,  # Yellow
        "red": 6,     # Red
        "blue": 3,    # Turquoise
        "pink": 5,    # Pink
        "gray": 16,   # Gray
    }

    def __init__(self, config: Optional[AnnotationConfig] = None):
        """
        Initialize the annotation manager.
        
        Args:
            config: Annotation configuration.
        """
        self.config = config or AnnotationConfig()
        self.annotations: List[Annotation] = []

    def create_annotation(self, modification: Modification) -> Annotation:
        """
        Create an annotation for a modification.
        
        Args:
            modification: The modification to annotate.
            
        Returns:
            Created Annotation object.
        """
        import uuid
        
        annotation = Annotation(
            id=str(uuid.uuid4()),
            modification_id=modification.id,
            source_ts_paragraph_id=modification.source_ts_paragraph_id,
            action_type=modification.action,
            confidence=modification.confidence,
            timestamp=datetime.utcnow().isoformat(),
            additional_info=modification.annotations.copy(),
        )
        
        self.annotations.append(annotation)
        return annotation

    def apply_annotation_to_run(
        self,
        run: Run,
        annotation: Annotation,
    ) -> None:
        """
        Apply annotation styling to a run.
        
        Args:
            run: The run to annotate.
            annotation: The annotation to apply.
        """
        # Apply highlight color based on action type
        highlight_color = self._get_highlight_color(annotation.action_type)
        if highlight_color:
            self._set_highlight_color(run, highlight_color)

    def add_inline_annotation(
        self,
        para: Paragraph,
        annotation: Annotation,
    ) -> Run:
        """
        Add inline annotation text to a paragraph.
        
        Args:
            para: The paragraph to annotate.
            annotation: The annotation to add.
            
        Returns:
            The created annotation run.
        """
        annotation_text = " " + annotation.to_text(self.config)
        
        run = para.add_run(annotation_text)
        run.font.size = Pt(self.config.font_size)
        run.font.color.rgb = RGBColor(128, 128, 128)
        run.italic = True
        
        return run

    def add_margin_annotation(
        self,
        para: Paragraph,
        annotation: Annotation,
    ) -> None:
        """
        Add margin annotation (Word comment) to a paragraph.
        
        Args:
            para: The paragraph to annotate.
            annotation: The annotation to add.
        """
        # Create comment text
        comment_text = self._build_comment_text(annotation)
        
        # Add Word comment using OOXML
        self._add_word_comment(para, comment_text, annotation.id)

    def _build_comment_text(self, annotation: Annotation) -> str:
        """Build comment text for margin annotation."""
        lines = [
            f"来源: TS 段落 {annotation.source_ts_paragraph_id}",
            f"操作: {annotation.action_type.value}",
            f"置信度: {annotation.confidence:.0%}",
            f"时间: {annotation.timestamp}",
        ]
        
        if annotation.additional_info:
            for key, value in annotation.additional_info.items():
                if key not in ("source_ts_paragraph_id", "action_type", "confidence_score"):
                    lines.append(f"{key}: {value}")
        
        return "\n".join(lines)

    def _add_word_comment(
        self,
        para: Paragraph,
        comment_text: str,
        comment_id: str,
    ) -> None:
        """
        Add a Word comment to a paragraph using OOXML.
        
        Args:
            para: The paragraph to comment on.
            comment_text: The comment text.
            comment_id: Unique identifier for the comment.
        """
        # Get or create comments part
        try:
            # Create comment range start
            comment_start = OxmlElement('w:commentRangeStart')
            comment_start.set(qn('w:id'), comment_id[:8])
            
            # Create comment range end
            comment_end = OxmlElement('w:commentRangeEnd')
            comment_end.set(qn('w:id'), comment_id[:8])
            
            # Create comment reference
            comment_ref = OxmlElement('w:commentReference')
            comment_ref.set(qn('w:id'), comment_id[:8])
            
            # Insert into paragraph
            para._p.insert(0, comment_start)
            para._p.append(comment_end)
            
            # Add comment reference in a run
            run = para.add_run()
            run._r.append(comment_ref)
            
        except Exception as e:
            logger.warning(f"Failed to add Word comment: {e}")
            # Fall back to inline annotation
            self.add_inline_annotation(para, Annotation(
                id=comment_id,
                modification_id="",
                source_ts_paragraph_id="",
                action_type=ActionType.INSERT,
                confidence=0.0,
                timestamp=datetime.utcnow().isoformat(),
            ))

    def _get_highlight_color(self, action_type: ActionType) -> Optional[int]:
        """Get highlight color code for action type."""
        if action_type == ActionType.INSERT and self.config.highlight_insertions:
            return self.HIGHLIGHT_COLORS.get(self.config.insertion_color)
        elif action_type == ActionType.OVERRIDE and self.config.highlight_overrides:
            return self.HIGHLIGHT_COLORS.get(self.config.override_color)
        return None

    def _set_highlight_color(self, run: Run, color_code: int) -> None:
        """Set highlight color on a run using OOXML."""
        try:
            # Access the run's XML element
            rPr = run._r.get_or_add_rPr()
            
            # Create highlight element
            highlight = OxmlElement('w:highlight')
            highlight.set(qn('w:val'), self._color_code_to_name(color_code))
            
            rPr.append(highlight)
        except Exception as e:
            logger.warning(f"Failed to set highlight color: {e}")

    def _color_code_to_name(self, color_code: int) -> str:
        """Convert color code to Word color name."""
        color_names = {
            3: "cyan",
            4: "green",
            5: "magenta",
            6: "red",
            7: "yellow",
            16: "darkGray",
        }
        return color_names.get(color_code, "yellow")

    def apply_annotations_to_document(
        self,
        doc: Document,
        modifications: List[Modification],
    ) -> None:
        """
        Apply annotations to all modifications in a document.
        
        Args:
            doc: The document to annotate.
            modifications: List of modifications to annotate.
        """
        for mod in modifications:
            annotation = self.create_annotation(mod)
            
            # Find the paragraph containing the modification
            for para in doc.paragraphs:
                if mod.new_text in para.text:
                    if self.config.style == AnnotationStyle.INLINE:
                        self.add_inline_annotation(para, annotation)
                    elif self.config.style == AnnotationStyle.MARGIN:
                        self.add_margin_annotation(para, annotation)
                    break

    def get_annotations_for_modification(
        self,
        modification_id: str,
    ) -> List[Annotation]:
        """Get all annotations for a specific modification."""
        return [a for a in self.annotations if a.modification_id == modification_id]

    def export_annotations_summary(self) -> List[Dict[str, Any]]:
        """Export all annotations as a summary list."""
        return [
            {
                "id": a.id,
                "modification_id": a.modification_id,
                "source_ts_paragraph_id": a.source_ts_paragraph_id,
                "action_type": a.action_type.value,
                "confidence": a.confidence,
                "timestamp": a.timestamp,
                **a.additional_info,
            }
            for a in self.annotations
        ]

    def clear_annotations(self) -> None:
        """Clear all stored annotations."""
        self.annotations = []
