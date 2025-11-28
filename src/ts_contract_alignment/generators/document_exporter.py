"""Document export functionality for generated contracts."""

import copy
import logging
import os
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from ..interfaces.generator import GeneratedContract, Modification
from ..models.enums import ActionType
from .annotation_manager import AnnotationConfig, AnnotationManager, AnnotationStyle


logger = logging.getLogger(__name__)


class DocumentExporter:
    """
    Exports generated contracts to .docx format.
    
    Supports revision-tracked and clean versions with
    full style and structure preservation.
    """

    def __init__(
        self,
        output_dir: str = "data/generated",
        annotation_config: Optional[AnnotationConfig] = None,
    ):
        """
        Initialize the document exporter.
        
        Args:
            output_dir: Directory for exported files.
            annotation_config: Configuration for annotations.
        """
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.annotation_config = annotation_config or AnnotationConfig()
        self.annotation_manager = AnnotationManager(self.annotation_config)

    def export_revision_tracked(
        self,
        contract: GeneratedContract,
        template_path: str,
    ) -> str:
        """
        Export contract with revision tracking marks.
        
        Args:
            contract: The generated contract.
            template_path: Path to the template document.
            
        Returns:
            Path to the exported file.
        """
        doc = Document(template_path)
        
        # Apply modifications with tracking
        self._apply_modifications_with_tracking(doc, contract.modifications)
        
        # Enable track changes mode
        self._enable_track_changes(doc)
        
        # Save the document
        output_path = contract.revision_tracked_path
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        
        logger.info(f"Exported revision-tracked document to: {output_path}")
        return output_path

    def export_clean_version(
        self,
        contract: GeneratedContract,
        template_path: str,
    ) -> str:
        """
        Export clean final version without revision marks.
        
        Args:
            contract: The generated contract.
            template_path: Path to the template document.
            
        Returns:
            Path to the exported file.
        """
        doc = Document(template_path)
        
        # Apply modifications without annotations
        self._apply_modifications_clean(doc, contract.modifications)
        
        # Save the document
        output_path = contract.clean_version_path
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        
        logger.info(f"Exported clean document to: {output_path}")
        return output_path

    def export_both_versions(
        self,
        contract: GeneratedContract,
        template_path: str,
    ) -> Tuple[str, str]:
        """
        Export both revision-tracked and clean versions.
        
        Args:
            contract: The generated contract.
            template_path: Path to the template document.
            
        Returns:
            Tuple of (revision_tracked_path, clean_version_path).
        """
        revision_path = self.export_revision_tracked(contract, template_path)
        clean_path = self.export_clean_version(contract, template_path)
        return revision_path, clean_path

    def _apply_modifications_with_tracking(
        self,
        doc: Document,
        modifications: List[Modification],
    ) -> None:
        """
        Apply modifications with revision tracking.
        
        Args:
            doc: The document to modify.
            modifications: List of modifications to apply.
        """
        # Sort modifications by location (reverse order)
        sorted_mods = sorted(
            modifications,
            key=lambda m: m.location_start,
            reverse=True,
        )
        
        for mod in sorted_mods:
            self._apply_tracked_modification(doc, mod)

    def _apply_tracked_modification(
        self,
        doc: Document,
        mod: Modification,
    ) -> None:
        """Apply a single modification with tracking."""
        # Find the target paragraph
        target_para = self._find_target_paragraph(doc, mod)
        
        if target_para is None:
            if mod.action == ActionType.INSERT:
                # Add new paragraph for insertions
                self._add_new_paragraph_with_tracking(doc, mod)
            return
        
        if mod.action == ActionType.OVERRIDE:
            self._apply_override_with_tracking(target_para, mod)
        elif mod.action == ActionType.INSERT:
            self._apply_insert_with_tracking(target_para, mod)

    def _find_target_paragraph(
        self,
        doc: Document,
        mod: Modification,
    ) -> Optional[any]:
        """Find the paragraph containing the modification target."""
        for para in doc.paragraphs:
            if mod.original_text and mod.original_text in para.text:
                return para
        return None

    def _apply_override_with_tracking(
        self,
        para: any,
        mod: Modification,
    ) -> None:
        """Apply an override modification with tracking."""
        # Store original formatting
        original_formatting = self._capture_paragraph_formatting(para)
        
        # Replace text
        new_text = para.text.replace(mod.original_text, mod.new_text)
        
        # Clear and rebuild paragraph
        para.clear()
        
        # Add deleted text (strikethrough)
        if mod.original_text:
            del_run = para.add_run(mod.original_text)
            del_run.font.strike = True
            del_run.font.color.rgb = RGBColor(255, 0, 0)
        
        # Add inserted text (highlighted)
        ins_run = para.add_run(mod.new_text)
        ins_run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
        
        # Apply original formatting to new run
        self._apply_formatting_to_run(ins_run, original_formatting)
        
        # Add annotation
        self._add_tracking_annotation(para, mod)

    def _apply_insert_with_tracking(
        self,
        para: any,
        mod: Modification,
    ) -> None:
        """Apply an insert modification with tracking."""
        # Add the new text with highlight
        ins_run = para.add_run(f" {mod.new_text}")
        ins_run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
        
        # Add annotation
        self._add_tracking_annotation(para, mod)

    def _add_new_paragraph_with_tracking(
        self,
        doc: Document,
        mod: Modification,
    ) -> None:
        """Add a new paragraph for inserted content."""
        para = doc.add_paragraph()
        
        # Add the new text with highlight
        ins_run = para.add_run(mod.new_text)
        ins_run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
        
        # Add annotation
        self._add_tracking_annotation(para, mod)

    def _add_tracking_annotation(
        self,
        para: any,
        mod: Modification,
    ) -> None:
        """Add tracking annotation to a paragraph."""
        annotation_text = f" 【TS:{mod.source_ts_paragraph_id} | {mod.action.value.upper()} | 置信度:{mod.confidence:.0%}】"
        
        ann_run = para.add_run(annotation_text)
        ann_run.font.size = Pt(8)
        ann_run.font.color.rgb = RGBColor(128, 128, 128)
        ann_run.italic = True

    def _apply_modifications_clean(
        self,
        doc: Document,
        modifications: List[Modification],
    ) -> None:
        """
        Apply modifications without tracking marks.
        
        Args:
            doc: The document to modify.
            modifications: List of modifications to apply.
        """
        # Sort modifications by location (reverse order)
        sorted_mods = sorted(
            modifications,
            key=lambda m: m.location_start,
            reverse=True,
        )
        
        for mod in sorted_mods:
            self._apply_clean_modification(doc, mod)

    def _apply_clean_modification(
        self,
        doc: Document,
        mod: Modification,
    ) -> None:
        """Apply a single modification without tracking."""
        target_para = self._find_target_paragraph(doc, mod)
        
        if target_para is None:
            if mod.action == ActionType.INSERT:
                para = doc.add_paragraph()
                para.add_run(mod.new_text)
            return
        
        if mod.action == ActionType.OVERRIDE:
            # Store original formatting
            original_formatting = self._capture_paragraph_formatting(target_para)
            
            # Replace text
            new_text = target_para.text.replace(mod.original_text, mod.new_text)
            
            # Clear and rebuild
            target_para.clear()
            run = target_para.add_run(new_text)
            self._apply_formatting_to_run(run, original_formatting)
            
        elif mod.action == ActionType.INSERT:
            target_para.add_run(f" {mod.new_text}")

    def _capture_paragraph_formatting(self, para: any) -> Dict:
        """Capture formatting from a paragraph."""
        formatting = {
            "alignment": para.alignment,
            "style": para.style.name if para.style else None,
        }
        
        if para.runs:
            first_run = para.runs[0]
            formatting["bold"] = first_run.bold
            formatting["italic"] = first_run.italic
            formatting["underline"] = first_run.underline
            formatting["font_name"] = first_run.font.name
            formatting["font_size"] = first_run.font.size
        
        return formatting

    def _apply_formatting_to_run(self, run: any, formatting: Dict) -> None:
        """Apply formatting to a run."""
        if formatting.get("bold") is not None:
            run.bold = formatting["bold"]
        if formatting.get("italic") is not None:
            run.italic = formatting["italic"]
        if formatting.get("underline") is not None:
            run.underline = formatting["underline"]
        if formatting.get("font_name"):
            run.font.name = formatting["font_name"]
        if formatting.get("font_size"):
            run.font.size = formatting["font_size"]

    def _enable_track_changes(self, doc: Document) -> None:
        """Enable track changes mode in the document."""
        try:
            # Access document settings
            settings = doc.settings.element
            
            # Create trackRevisions element
            track_revisions = OxmlElement('w:trackRevisions')
            settings.append(track_revisions)
            
        except Exception as e:
            logger.warning(f"Could not enable track changes: {e}")

    def generate_diff_report(
        self,
        contract: GeneratedContract,
    ) -> str:
        """
        Generate a text-based diff report of all modifications.
        
        Args:
            contract: The generated contract.
            
        Returns:
            Diff report as a string.
        """
        lines = [
            "=" * 60,
            "CONTRACT MODIFICATION REPORT",
            f"Generated: {contract.generation_timestamp}",
            f"Template ID: {contract.template_document_id}",
            f"TS Document ID: {contract.ts_document_id}",
            "=" * 60,
            "",
        ]
        
        for i, mod in enumerate(contract.modifications, 1):
            lines.extend([
                f"Modification #{i}",
                f"  ID: {mod.id}",
                f"  Action: {mod.action.value}",
                f"  Source TS Paragraph: {mod.source_ts_paragraph_id}",
                f"  Confidence: {mod.confidence:.2%}",
                f"  Original: {mod.original_text[:100]}..." if len(mod.original_text) > 100 else f"  Original: {mod.original_text}",
                f"  New: {mod.new_text[:100]}..." if len(mod.new_text) > 100 else f"  New: {mod.new_text}",
                "",
            ])
        
        lines.extend([
            "=" * 60,
            f"Total Modifications: {len(contract.modifications)}",
            f"Insertions: {sum(1 for m in contract.modifications if m.action == ActionType.INSERT)}",
            f"Overrides: {sum(1 for m in contract.modifications if m.action == ActionType.OVERRIDE)}",
            "=" * 60,
        ])
        
        return "\n".join(lines)

    def export_diff_report(
        self,
        contract: GeneratedContract,
        output_path: Optional[str] = None,
    ) -> str:
        """
        Export diff report to a file.
        
        Args:
            contract: The generated contract.
            output_path: Optional output path.
            
        Returns:
            Path to the exported report.
        """
        if output_path is None:
            output_path = str(
                self.output_dir / f"diff_report_{contract.id[:8]}.txt"
            )
        
        report = self.generate_diff_report(contract)
        
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(report)
        
        logger.info(f"Exported diff report to: {output_path}")
        return output_path
